# main.py
import os
import re
import json
import time
import asyncio
import hashlib
import pandas as pd
import pdfplumber
from typing import Optional, Dict, Any, List, Tuple
from datetime import datetime
from fastapi import FastAPI, UploadFile, File, BackgroundTasks
from pydantic import BaseModel
from dotenv import load_dotenv

from langchain_openai import ChatOpenAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from core import (
    analyze_dataframe, generate_schema_description,
    validate_pandas_code, safe_execute_pandas,
    get_pandas_prompt, get_natural_language_prompt,
    preprocess_query, get_fuzzy_column_matches, extract_code
)

# ==================== CONFIG ====================
load_dotenv()

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.normpath(os.path.join(BASE_DIR, '..'))

DATA_DIR = os.path.join(ROOT_DIR, 'data')
VECTOR_STORE_DIR = os.path.join(ROOT_DIR, 'vector_store')
METRICS_DIR = os.path.join(ROOT_DIR, 'metrics')
HASH_STORE_FILE = os.path.join(VECTOR_STORE_DIR, 'content_hashes.json')

for d in [DATA_DIR, VECTOR_STORE_DIR, METRICS_DIR]: os.makedirs(d, exist_ok=True)

app = FastAPI()
structured_data_store = {}

# ==================== METRICS TRACKER ====================
class MetricsTracker:
    def __init__(self, query: str):
        self.query = query
        self.start_time = time.time()
        self.metrics = {"query": query, "timestamp": datetime.now().isoformat(), "timing": {}, "rag_details": {}, "routing": {}, "rag_triad_scores": {}}
    
    def start(self, phase: str): self.metrics["timing"][phase] = {"start": time.time()}
    def end(self, phase: str):
        if phase in self.metrics["timing"]:
            self.metrics["timing"][phase]["elapsed_ms"] = round((time.time() - self.metrics["timing"][phase]["start"]) * 1000, 2)
    def set_rag_details(self, key: str, value: Any): self.metrics["rag_details"][key] = value
    def set_routing(self, key: str, value: Any): self.metrics["routing"][key] = value

    def calculate_rag_triad(self, query: str, context: str, answer: str, llm_client):
        try:
            prompt = f"""Score 1-10 JSON only: {{"context_relevance": 8, "groundedness": 9, "answer_relevance": 7}}\nQuery:{query}\nContext:{str(context)[:500]}\nAnswer:{answer}"""
            scores = llm_client.invoke(prompt).content.replace("```json","").replace("```","").strip()
            self.metrics["rag_triad_scores"] = json.loads(scores)
        except: self.metrics["rag_triad_scores"] = {"error": "eval_failed"}

    def finalize(self, success: bool):
        self.metrics["timing"]["total"] = {"elapsed_ms": round((time.time() - self.start_time) * 1000, 2)}
        self.metrics["success"] = success
        try:
            date_str = datetime.now().strftime("%Y-%m-%d")
            with open(os.path.join(METRICS_DIR, f"rag_metrics_{date_str}.jsonl"), 'a', encoding='utf-8') as f:
                f.write(json.dumps(self.metrics, ensure_ascii=False, default=str) + '\n')
        except: pass

# ==================== UTILITIES ====================
def sanitize_filename(filename: str) -> str:
    _, ext = os.path.splitext(filename)
    return f"uploaded_{int(time.time())}{ext}"

# ==================== CONTENT HASH STORE ====================
class ContentHashStore:
    def __init__(self, hash_file: str):
        self.hash_file = hash_file
        self.hashes = self._load()
    def _load(self) -> Dict[str, Dict]:
        if os.path.exists(self.hash_file):
            try:
                with open(self.hash_file, 'r', encoding='utf-8') as f: return json.load(f)
            except: return {}
        return {}
    def _save(self):
        with open(self.hash_file, 'w', encoding='utf-8') as f: json.dump(self.hashes, f, indent=2)
    def get_hash(self, content: str) -> str: return hashlib.sha256(content.encode('utf-8')).hexdigest()
    def exists(self, content: str) -> Tuple[bool, Optional[str]]:
        h = self.get_hash(content)
        return (True, self.hashes[h].get("filename")) if h in self.hashes else (False, None)
    def add(self, content: str, filename: str, doc_type: str):
        self.hashes[self.get_hash(content)] = {"filename": filename, "type": doc_type}
        self._save()

hash_store = ContentHashStore(HASH_STORE_FILE)

# ==================== VECTOR STORES ====================
print("🔧 Initializing Embeddings Model...")
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
doc_vector_store = Chroma(persist_directory=os.path.join(VECTOR_STORE_DIR, 'docs'), embedding_function=embeddings)
schema_vector_store = Chroma(persist_directory=os.path.join(VECTOR_STORE_DIR, 'schemas'), embedding_function=embeddings)
print("✅ Embeddings and Vector Stores Loaded.")

# ==================== LLM ====================
llm = ChatOpenAI(
    model="meta/llama-3.1-8b-instruct",
    openai_api_key=os.getenv("NVIDIA_API_KEY"),
    openai_api_base="https://integrate.api.nvidia.com/v1",
    temperature=0
)

# ==================== INGESTION ====================
def process_structured(path: str, filename: str, skip_dedup: bool = False) -> Dict[str, Any]:
    print(f"  📂 Reading structured file: {filename}")
    try:
        df = pd.read_csv(path) if filename.endswith('.csv') else pd.read_excel(path)
    except Exception as e: 
        print(f"  ❌ FAILED to read {filename}: {str(e)}")
        return {"success": False, "message": f"Read error: {str(e)}"}
    
    df.columns = df.columns.str.strip().str.lower().str.replace(" ", "_")
    original_len = len(df)
    if original_len > 1000: df = df.head(1000)
    
    file_content = df.to_csv(index=False)
    if not skip_dedup:
        exists, existing_file = hash_store.exists(file_content)
        if exists:
            print(f"  ⏭️  SKIPPED {filename} (Duplicate of {existing_file}). Loading 1000 rows to RAM.")
            metadata = analyze_dataframe(df)
            structured_data_store[filename] = {"df": df, "metadata": metadata}
            return {"skipped": True, "message": f"Duplicate"}
    
    metadata = analyze_dataframe(df)
    structured_data_store[filename] = {"df": df, "metadata": metadata}
    schema_vector_store.add_documents([Document(page_content=generate_schema_description(filename, df, metadata), metadata={"filename": filename, "type": "schema"})])
    hash_store.add(file_content, filename, "structured")
    
    print(f"  ✅ LOADED {filename} ({len(df)} rows, {len(df.columns)} cols) -> Added to Vector DB & RAM")
    return {"success": True, "message": f"Loaded: {filename}", "rows": len(df), "columns": len(df.columns)}

def process_unstructured(path: str, filename: str, skip_dedup: bool = False) -> Dict[str, Any]:
    print(f"  📂 Reading unstructured file: {filename}")
    text = ""
    try:
        if filename.endswith('.pdf'):
            text_blocks = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    if page.extract_text(): text_blocks.append(page.extract_text())
                    for table in page.extract_tables():
                        markdown_table = "\n".join([" | ".join([str(cell or "") for cell in row]) for row in table])
                        text_blocks.append(f"--- FINANCIAL TABLE ---\n{markdown_table}\n--- END TABLE ---")
            text = "\n\n".join(text_blocks)
        elif filename.endswith('.docx'):
            try:
                import docx
                doc = docx.Document(path)
                text = "\n".join([p.text for p in doc.paragraphs])
            except: text = ""
        else:
            with open(path, encoding="utf-8", errors="ignore") as f: text = f.read()
    except Exception as e: 
        print(f"  ❌ FAILED to read {filename}: {str(e)}")
        return {"success": False, "message": str(e)}
    
    if not text.strip(): 
        print(f"  ⚠️  SKIPPED {filename} (No text extracted)")
        return {"success": False, "message": "No text extracted"}
    
    if not skip_dedup:
        exists, existing_file = hash_store.exists(text)
        if exists: 
            print(f"  ⏭️  SKIPPED {filename} (Duplicate of {existing_file})")
            return {"skipped": True, "message": f"Duplicate"}
    
    splits = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200).split_documents([Document(page_content=text, metadata={"source": filename})])
    for split in splits:
        year_match = re.search(r'\b(20\d{2})\b', split.page_content)
        split.metadata["year"] = year_match.group(1) if year_match else "unknown"

    doc_vector_store.add_documents(splits)
    hash_store.add(text, filename, "unstructured")
    print(f"  ✅ LOADED {filename} ({len(splits)} chunks) -> Added to Vector DB")
    return {"success": True, "message": f"Loaded: {filename}", "chunks": len(splits)}

# ==================== STARTUP & UPLOAD ====================
@app.on_event("startup")
def startup():
    print("\n" + "="*70)
    print("🚀 STARTING UP SYSTEM (Auto-Restart Safe)")
    print("="*70)
    print(f"📂 Looking for files in: {DATA_DIR}")
    
    if not os.path.exists(DATA_DIR):
        print("  ⚠️  DATA DIRECTORY DOES NOT EXIST!")
        return
    if not os.listdir(DATA_DIR):
        print("  ⚠️  DATA DIRECTORY IS EMPTY!")
        return
    
    # CHECK: Force reload if vector stores are empty but hash store has entries
    force_reload = False
    try:
        schema_count = schema_vector_store._collection.count()
        doc_count = doc_vector_store._collection.count()
        hash_count = len(hash_store.hashes)
        
        if hash_count > 0 and (schema_count == 0 or doc_count == 0):
            print("  🔄 Detected mismatch. Force reloading all files...")
            force_reload = True
            hash_store.hashes = {}
            hash_store._save()
    except Exception as e:
        print(f"  ⚠️ Could not check vector store counts: {e}. Force reloading...")
        force_reload = True
        hash_store.hashes = {}
        hash_store._save()
    
    for file in os.listdir(DATA_DIR):
        path = os.path.join(DATA_DIR, file)
        if not os.path.isfile(path): continue
        if file.endswith(('.csv', '.xlsx')): 
            process_structured(path, file, skip_dedup=False)
        elif file.endswith(('.pdf', '.txt', '.docx')): 
            process_unstructured(path, file, skip_dedup=False)
        else: 
            print(f"  ⏭️  SKIPPED {file} (Unsupported format)")
            
    print(f"💾 Total Datasets Ready in RAM: {len(structured_data_store)}")
    print("="*70 + "\n")

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    safe_name = sanitize_filename(file.filename)
    file_ext = os.path.splitext(safe_name)[1].lower()
    if file_ext not in ['.csv', '.xlsx', '.pdf', '.txt', '.docx']: return {"success": False, "message": f"Unsupported: {file_ext}"}
    
    contents = await file.read()
    if len(contents) > 10 * 1024 * 1024: return {"success": False, "message": "File too large (Max 10MB)"}
    
    file_path = os.path.join(DATA_DIR, safe_name)
    with open(file_path, "wb") as f: f.write(contents)
    return process_structured(file_path, safe_name) if file_ext in ['.csv', '.xlsx'] else process_unstructured(file_path, safe_name)

# ==================== CORE LOGIC ====================
async def get_candidates(query: str) -> List[Dict]:
    candidates = []
    tokens = query.lower().split()
    try: schema_hits = await asyncio.to_thread(schema_vector_store.similarity_search, query, k=3)
    except: schema_hits = []
    
    schema_scores = {}
    for i, hit in enumerate(schema_hits):
        fname = hit.metadata.get("filename")
        if fname: schema_scores[fname] = schema_scores.get(fname, 0) + (3 - i) * 5
    
    for filename, store in structured_data_store.items():
        df, metadata = store["df"], store["metadata"]
        score = schema_scores.get(filename, 0)
        matched_columns = get_fuzzy_column_matches(tokens, list(df.columns))
        nums = re.findall(r"\d+", query)
        for col in df.columns:
            for num in nums:
                if df[col].astype(str).eq(num).any():
                    score += 10
                    if col == metadata.get("primary_key"): score += 30
        score += len(matched_columns) * 10
        if score > 0: candidates.append({"filename": filename, "score": score, "matched_columns": matched_columns})
    candidates.sort(key=lambda x: x["score"], reverse=True)
    return candidates

async def classify_intent(query: str, candidates: List[Dict], metrics: MetricsTracker) -> Dict[str, Any]:
    metrics.start("intent_classification")
    
    # PRE-CHECK: Force concept_explain for system/meta queries without asking LLM
    system_keywords = ["system", "technologies", "technology", "architecture", "workflow", "purpose of this", "how does this work", "role of"]
    if any(kw in query.lower() for kw in system_keywords):
        metrics.end("intent_classification")
        print(f"  🧠 LLM Intent: {{'intent': 'concept_explain', 'is_data_query': False, 'target_dataset': None}} (Keyword Pre-check)")
        return {"intent": "concept_explain", "is_data_query": False, "target_dataset": None}

    cand_info = "\n".join([f"{i+1}. {c['filename']} (Score: {c['score']})" for i, c in enumerate(candidates[:3])]) if candidates else "None"
    
    prompt = f"""QUERY: "{query}"
DATASETS:
{cand_info}

You must respond with ONLY a single JSON object, nothing else.
Valid intents: "data_lookup" or "concept_explain"
If the query asks about the datasets, use "data_lookup". If it asks about concepts, definitions, or the system itself, use "concept_explain".

Example: {{"intent": "data_lookup", "is_data_query": true, "target_dataset": "filename.csv"}}

YOUR JSON:"""
    
    try:
        resp_text = (await llm.ainvoke(prompt)).content.strip()
        
        # FIX: Remove newline characters that break JSON parsing
        resp_text = resp_text.replace("\n", " ").replace("\r", " ")
        
        resp_text = resp_text.replace("```json", "").replace("```", "").strip()
        
        json_match = re.search(r'\{[^{}]*\}', resp_text, re.DOTALL)
        
        if json_match:
            resp = json.loads(json_match.group())
            if "intent" not in resp: resp["intent"] = "unknown"
            if "is_data_query" not in resp: resp["is_data_query"] = False
            if "target_dataset" not in resp: resp["target_dataset"] = None
        else:
            resp = {"intent": "unknown", "is_data_query": False, "target_dataset": None}
        
        metrics.end("intent_classification")
        print(f"  🧠 LLM Intent: {resp}")
        return resp
        
    except json.JSONDecodeError as e:
        metrics.end("intent_classification")
        print(f"  ⚠️ LLM Intent Failed to Parse: {str(e)}")
        return {"intent": "unknown", "is_data_query": False, "target_dataset": None}
    except Exception as e:
        metrics.end("intent_classification")
        print(f"  ⚠️ LLM Intent Error: {str(e)}")
        return {"intent": "unknown", "is_data_query": False, "target_dataset": None}

async def execute_structured(filename: str, query: str, metrics: MetricsTracker) -> Dict[str, Any]:
    metrics.start("structured_execution")
    print(f"  🔍 Executing Pandas on: {filename}")
    df = structured_data_store[filename]["df"]
    metadata = structured_data_store[filename]["metadata"]
    raw_data_sample = df.head(3).to_string()
    
    last_error = None  # Track actual error
    
    for attempt in range(2):
        metrics.start("code_generation")
        prompt = get_pandas_prompt(query, metadata, raw_data_sample, last_error)
        code = extract_code((await llm.ainvoke(prompt)).content)
        metrics.end("code_generation")
        print(f"  🐍 Pandas Code Gen (Attempt {attempt+1}): {code[:100]}...")
        
        valid, msg = validate_pandas_code(code, list(df.columns), metadata)
        if not valid: 
            print(f"  ❌ Code Validation Failed: {msg}")
            last_error = f"VALIDATION ERROR: {msg}. Do NOT use imports."
            continue
        
        metrics.start("code_execution")
        result, err = safe_execute_pandas(code, df)
        metrics.end("code_execution")
        
        if not err:
            metrics.end("structured_execution")
            print(f"  ✅ Pandas Execution Success.")
            return {"type": "structured", "context": str(result), "dataset_used": filename}
        else:
            print(f"  ❌ Pandas Execution Error: {err}")
            last_error = err  # Pass exact error to next attempt
        
    metrics.end("structured_execution")
    return {"type": "error", "context": "execution_failed"}

async def execute_unstructured(query: str, metrics: MetricsTracker) -> Dict[str, Any]:
    metrics.start("unstructured_search")
    query_year_match = re.search(r'\b(20\d{2})\b', query)
    where_filter = {"year": query_year_match.group(1)} if query_year_match else None
    
    try: results = await asyncio.to_thread(doc_vector_store.similarity_search, query, k=3, filter=where_filter)
    except: results = []
    
    metrics.end("unstructured_search")
    if not results: 
        print(f"  ⚠️ Unstructured Search: No PDF chunks found.")
        return {"type": "none", "context": ""}
    print(f"  📄 Unstructured Search: Found {len(results)} chunks.")
    return {"type": "unstructured", "context": "\n\n---\n\n".join([d.page_content for d in results])}

async def route_and_execute(query: str) -> Dict[str, Any]:
    metrics = MetricsTracker(query)
    clean_query = preprocess_query(query)
    if clean_query == "Invalid query detected.": return {"type": "none", "context": "", "metrics": metrics}
    
    print(f"\n{'='*50}")
    print(f"🗣️  USER QUERY: {query}")
    print(f"{'='*50}")
    
    candidates = await get_candidates(clean_query)
    print(f"📊 Candidates Found: {len(candidates)} - {[c['filename'] + ' (Score:'+str(c['score'])+')' for c in candidates[:3]]}")
    
    intent = await classify_intent(clean_query, candidates, metrics)
    
    is_data_query = intent.get("is_data_query", False)
    target_dataset = intent.get("target_dataset")
    intent_type = intent.get("intent", "unknown")
    
    # SCENARIO 1: DATA LOOKUP
    if is_data_query:
        print(f"🛤️  ROUTE: DATA LOOKUP path")
        if target_dataset and target_dataset in structured_data_store:
            print(f"   → Trying: {target_dataset} (exact match)")
            res = await execute_structured(target_dataset, clean_query, metrics)
            if res["type"] != "error": return {**res, "metrics": metrics}
        
        if candidates and candidates[0]["score"] >= 30:
            print(f"   → Trying: {candidates[0]['filename']} (score fallback)")
            res = await execute_structured(candidates[0]["filename"], clean_query, metrics)
            if res["type"] != "error": return {**res, "metrics": metrics}
        
        print(f"   ❌ Structured execution failed. NOT falling back to unstructured (safety)")
        metrics.set_routing("decision", "data_query_failed")
        return {"type": "none", "context": "", "metrics": metrics}
    
    # SCENARIO 2: CONCEPT EXPLAIN
    elif intent_type == "concept_explain":
        print(f"🛤️  ROUTE: CONCEPT EXPLAIN path -> Unstructured Docs")
        res = await execute_unstructured(clean_query, metrics)
        return {**res, "metrics": metrics}
    
    # SCENARIO 3: UNKNOWN
    else:
        print(f"🛤️  ROUTE: UNKNOWN intent -> Trying Unstructured as last resort")
        res = await execute_unstructured(clean_query, metrics)
        if res["type"] != "none": return {**res, "metrics": metrics}
        
        print(f"🛤️  ROUTE: NO MATCH FOUND anywhere")
        metrics.set_routing("decision", "no_match")
        return {"type": "none", "context": "", "metrics": metrics}

# ==================== ENDPOINTS ====================
class QueryRequest(BaseModel):
    question: str

@app.post("/query")
async def query(req: QueryRequest, background_tasks: BackgroundTasks):
    result = await route_and_execute(req.question)
    metrics = result.get("metrics", MetricsTracker(req.question))
    
    if result["type"] == "none" or result["type"] == "error":
        metrics.finalize(False)
        
        routing_decision = metrics.metrics.get("routing", {}).get("decision", "no_match")
        if routing_decision == "data_query_failed":
            warning_msg = "[WARNING] Could not retrieve the data. The dataset may not contain this information, or there was a processing error."
        else:
            warning_msg = "[WARNING] No relevant data found for this query."
            
        return {
            "answer": warning_msg,
            "execution_type": "none",
            "dataset_used": None,
            "generated_context": ""
        }
    
    answer = (await llm.ainvoke(get_natural_language_prompt(req.question, result["context"]))).content
    
    background_tasks.add_task(metrics.calculate_rag_triad, req.question, result["context"], answer, llm)
    background_tasks.add_task(metrics.finalize, True)
    
    return {
        "answer": answer,
        "execution_type": result["type"],
        "dataset_used": result.get("dataset_used"),
        "generated_context": result["context"]
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)